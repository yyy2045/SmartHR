"""
文档处理器 - PDF/Word 处理、分块和向量化
"""

import io
from typing import List, Dict, Any, Optional

from src.services.rag.chunker import semantic_chunk_text


class DocumentProcessor:
    """处理文档：提取文本、分块、向量化到知识库"""

    def __init__(self):
        self.chunk_size = 500
        self.chunk_overlap = 50

    async def process_document(self, file_path: str, metadata: Dict[str, Any]) -> List[str]:
        """
        处理文档文件：提取文本、分块、向量化
        返回存储在向量数据库中的 chunk ID 列表
        """
        # 从文件提取文本
        text = self.extract_text_from_file(file_path)

        # 对文本分块
        chunks = self.chunk_text(text)

        # 向量化并存储
        chunk_ids = await self.vectorize_chunks(chunks, metadata)

        return chunk_ids

    async def process_file_content(self, content: bytes, filename: str, metadata: Dict[str, Any]) -> List[str]:
        """
        直接处理上传的文件内容，不保存到磁盘
        返回 chunk ID 列表
        """
        # 根据文件类型提取文本
        text = self.extract_text_from_bytes(content, filename)

        # 对文本分块
        chunks = self.chunk_text(text)

        # 向量化并存储
        chunk_ids = await self.vectorize_chunks(chunks, metadata)

        return chunk_ids

    def extract_text_from_file(self, file_path: str) -> str:
        """从 PDF、Word 或 TXT 文件提取文本"""
        import os

        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._extract_word(file_path)
        elif ext in ['.txt', '.md']:
            return self._extract_text_file(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")

    def extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        """从文件字节提取文本"""
        ext = filename.lower().split('.')[-1]

        if ext == 'pdf':
            return self._extract_pdf_from_bytes(content)
        elif ext in ['docx', 'doc']:
            return self._extract_word_from_bytes(content)
        elif ext in ['txt', 'md']:
            return self._decode_text_bytes(content)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")

    def _extract_pdf(self, file_path: str) -> str:
        """从 PDF 文件提取文本"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("需要安装 PyPDF2 来处理 PDF。安装命令: pip install PyPDF2")

    def _extract_pdf_from_bytes(self, content: bytes) -> str:
        """从 PDF 字节提取文本"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("需要安装 PyPDF2 来处理 PDF")

    def _extract_word(self, file_path: str) -> str:
        """从 Word 文件提取文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("需要安装 python-docx 来处理 Word 文档。安装命令: pip install python-docx")

    def _extract_word_from_bytes(self, content: bytes) -> str:
        """从 Word 字节提取文本"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("需要安装 python-docx 来处理 Word 文档")

    def _extract_text_file(self, file_path: str) -> str:
        """从纯文本文件提取文本。"""
        with open(file_path, "rb") as file:
            return self._decode_text_bytes(file.read())

    def _decode_text_bytes(self, content: bytes) -> str:
        """兼容 UTF-8 和常见中文 Windows 编码。"""
        for encoding in ["utf-8-sig", "utf-8", "gb18030"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")

    def chunk_text(self, text: str, chunk_size: Optional[int] = None, overlap: Optional[int] = None) -> List[str]:
        """按段落、标题和句子优先切分，超长文本再固定长度兜底。"""
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap
        return semantic_chunk_text(text, chunk_size=chunk_size, overlap=overlap)

    async def vectorize_chunks(self, chunks: List[str], metadata: Dict[str, Any]) -> List[str]:
        """将文本块向量化并存储到 Chroma"""
        from src.services.rag.pipeline import rag_pipeline
        from src.services.rag.schemas import RagIndexRequest

        if not chunks:
            return []
        response = await rag_pipeline.index(
            RagIndexRequest(
                companyId=str(metadata.get("company_id") or metadata.get("companyId") or "default"),
                sourceType=str(metadata.get("source_type") or "knowledge"),
                sourceId=str(metadata.get("doc_id") or metadata.get("source_id") or "doc"),
                title=str(metadata.get("title") or metadata.get("filename") or ""),
                chunks=chunks,
                collection=str(metadata.get("collection") or "knowledge_base"),
                metadata=metadata,
            )
        )
        return response.chunkIds

    async def process_multiple_documents(self, files: List[Dict[str, Any]]) -> Dict[str, List[str]]:
        """
        处理多个文档
        files: {path, metadata} 字典列表
        返回: {filename: [chunk_ids]}
        """
        results = {}
        for file_info in files:
            path = file_info.get('path')
            metadata = file_info.get('metadata', {})
            try:
                chunk_ids = await self.process_document(path, metadata)
                results[path] = chunk_ids
            except Exception as e:
                results[path] = []
                print(f"处理 {path} 时出错: {e}")
        return results


# 全局实例
document_processor = DocumentProcessor()
